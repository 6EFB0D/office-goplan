# -*- coding: utf-8 -*-
"""書類生成モジュール。python-docx で Word 文書を生成し、win32com で PDF に変換する。"""
from __future__ import annotations

import datetime
import os
from dataclasses import dataclass, field
from enum import Enum

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

import config

# ---- デザイン定数 ----
COLOR_ACCENT  = RGBColor(0x00, 0xB8, 0x55)
COLOR_TEXT    = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_MUTED   = RGBColor(0x6B, 0x72, 0x80)
COLOR_DIVIDER = "D0D5DD"
FONT_JP       = "Noto Sans JP"


# ---- データクラス ----

class DocType(Enum):
    QUOTE   = "quote"
    INVOICE = "invoice"
    RECEIPT = "receipt"


@dataclass
class LineItem:
    description: str
    qty:         int
    unit_price:  int

    @property
    def amount(self) -> int:
        return self.qty * self.unit_price


@dataclass
class DocumentParams:
    doc_type:         DocType
    doc_number:       str
    issue_date:       datetime.date
    customer_name:    str
    customer_address: str
    items:            list[LineItem]
    due_date:         datetime.date | None = None
    valid_until:      datetime.date | None = None
    notes_extra:      list[str] = field(default_factory=list)

    @property
    def total(self) -> int:
        return sum(i.amount for i in self.items)


# ---- XML ヘルパー ----

def _set_cell_border(cell, **edges) -> None:
    tcp = cell._tc.get_or_add_tcPr()
    borders = tcp.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcp.append(borders)
    for name, attrs in edges.items():
        tag = qn(f"w:{name}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{name}")
            borders.append(el)
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), str(v))


def _shade_cell(cell, hex_color: str) -> None:
    tcp = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcp.append(shd)


def _apply_font(run, size_pt=None, bold=False, color=None) -> None:
    run.font.name = FONT_JP
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(attr), FONT_JP)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _para(container, text="", *, size_pt=None, bold=False, color=None,
          align=None, space_after_pt=None):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    if text:
        _apply_font(p.add_run(text), size_pt=size_pt, bold=bold, color=color)
    return p


def _nil():
    return {"val": "nil"}

def _single(color=COLOR_DIVIDER, sz="6"):
    return {"val": "single", "sz": sz, "color": color}


# ---- ロゴ準備 ----

def _prepare_logo() -> None:
    if os.path.exists(config.LOGO_RENDERED):
        try:
            if os.path.getmtime(config.LOGO_RENDERED) >= os.path.getmtime(config.LOGO_PATH):
                return
        except OSError:
            pass
    if not os.path.exists(config.LOGO_PATH):
        return
    try:
        from PIL import Image
    except ImportError:
        return
    with Image.open(config.LOGO_PATH) as img:
        img.save(config.LOGO_RENDERED, format="PNG", optimize=True)


# ---- ブロック生成 ----

def _setup_page(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_height = Mm(297); sec.page_width = Mm(210)
    sec.top_margin = Cm(1.5);  sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = FONT_JP
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rf.set(qn(attr), FONT_JP)


def _add_header_block(doc, title, meta_rows) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(11.0)
    tbl.columns[1].width = Cm(6.0)
    left  = tbl.rows[0].cells[0]; left.width  = Cm(11.0)
    right = tbl.rows[0].cells[1]; right.width = Cm(6.0)

    tp = left.paragraphs[0]
    tp.paragraph_format.space_after = Pt(6)
    _apply_font(tp.add_run(title), size_pt=26, bold=True, color=COLOR_TEXT)

    ap = left.add_paragraph()
    ap.paragraph_format.space_after = Pt(10)
    _apply_font(ap.add_run("▬▬▬▬"), size_pt=10, color=COLOR_ACCENT, bold=True)

    for label, value in meta_rows:
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _apply_font(p.add_run(f"{label}　"), size_pt=9, color=COLOR_MUTED)
        _apply_font(p.add_run(value), size_pt=10, color=COLOR_TEXT)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(config.LOGO_RENDERED):
        rp.add_run().add_picture(config.LOGO_RENDERED, width=Cm(5.0))
    else:
        _apply_font(rp.add_run("[Office Go Plan]"), size_pt=10, color=COLOR_MUTED)

    for cell in (left, right):
        _set_cell_border(cell,
            top=_nil(), bottom=_nil(), left=_nil(), right=_nil())


def _add_parties_block(doc, buyer_name, buyer_address="") -> None:
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(9.0)
    tbl.columns[1].width = Cm(8.0)
    left  = tbl.rows[0].cells[0]; left.width  = Cm(9.0)
    right = tbl.rows[0].cells[1]; right.width = Cm(8.0)

    p = left.paragraphs[0]; p.paragraph_format.space_after = Pt(4)
    _apply_font(p.add_run("請求先"), size_pt=9, color=COLOR_MUTED)
    p = left.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    _apply_font(p.add_run(f"{buyer_name}　御中"), size_pt=12, bold=True, color=COLOR_TEXT)
    if buyer_address:
        p = left.add_paragraph()
        _apply_font(p.add_run(buyer_address), size_pt=9, color=COLOR_TEXT)

    p = right.paragraphs[0]; p.paragraph_format.space_after = Pt(4)
    _apply_font(p.add_run("発行元"), size_pt=9, color=COLOR_MUTED)
    for txt, bold in [
        (config.SELLER_NAME, True),
        (f"{config.SELLER_POSTAL} {config.SELLER_ADDRESS}", False),
        (f"Email: {config.SELLER_EMAIL}", False),
        (config.SELLER_WEB, False),
    ]:
        p = right.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        _apply_font(p.add_run(txt), size_pt=10 if bold else 9, bold=bold, color=COLOR_TEXT)

    for cell in (left, right):
        _set_cell_border(cell,
            top=_nil(), bottom=_nil(), left=_nil(), right=_nil())


def _add_headline_amount(doc, label, amount_text, sub_note=None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(0)
    _para(doc, "", size_pt=2)

    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_border(cell,
        top=_single(), bottom=_single(), left=_nil(), right=_nil())
    cell.paragraphs[0].paragraph_format.space_before = Pt(6)
    cell.paragraphs[0].paragraph_format.space_after  = Pt(6)
    _apply_font(cell.paragraphs[0].add_run(f"{label}　"), size_pt=11, color=COLOR_MUTED)
    _apply_font(cell.paragraphs[0].add_run(amount_text),  size_pt=22, bold=True, color=COLOR_TEXT)
    if sub_note:
        p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        _apply_font(p.add_run(sub_note), size_pt=9, color=COLOR_MUTED)


def _add_line_items_table(doc, items: list[LineItem]) -> None:
    _para(doc, "", size_pt=2)
    widths = [Cm(9.0), Cm(2.0), Cm(3.0), Cm(3.0)]
    tbl = doc.add_table(rows=1 + len(items), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    for i, w in enumerate(widths):
        tbl.columns[i].width = w

    for i, h in enumerate(["説明", "数量", "単価（円）", "金額（円）"]):
        c = tbl.rows[0].cells[i]; c.width = widths[i]
        _shade_cell(c, "F6F8FA")
        _set_cell_border(c, top=_single(), bottom=_single(), left=_nil(), right=_nil())
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        if i >= 1: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _apply_font(p.add_run(h), size_pt=9, bold=True, color=COLOR_MUTED)

    for idx, item in enumerate(items, 1):
        row = tbl.rows[idx].cells
        for i, v in enumerate([item.description, f"{item.qty:,}",
                                f"{item.unit_price:,}", f"{item.amount:,}"]):
            row[i].width = widths[i]
            _set_cell_border(row[i],
                top=_nil(), bottom=_single(sz="4"), left=_nil(), right=_nil())
            p = row[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            if i >= 1: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _apply_font(p.add_run(v), size_pt=10, color=COLOR_TEXT)


def _add_totals_block(doc, total: int, total_label: str) -> None:
    _para(doc, "", size_pt=2)
    tbl = doc.add_table(rows=3, cols=2)
    tbl.autofit = False; tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    tbl.columns[0].width = Cm(11.0); tbl.columns[1].width = Cm(6.0)

    for i, (lbl, val, emph) in enumerate([
        ("小計",      f"{total:,}",   False),
        ("合計",      f"{total:,}",   False),
        (total_label, f"¥{total:,}", True),
    ]):
        left  = tbl.rows[i].cells[0]; left.width  = Cm(11.0)
        right = tbl.rows[i].cells[1]; right.width = Cm(6.0)
        _set_cell_border(left,  top=_nil(), bottom=_nil(), left=_nil(), right=_nil())
        _set_cell_border(right,
            top=_single() if emph else _nil(),
            bottom=_nil(), left=_nil(), right=_nil())
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
        _apply_font(p.add_run(f"{lbl}　"), size_pt=10, color=COLOR_MUTED)
        _apply_font(p.add_run(val), size_pt=13 if emph else 10, bold=emph, color=COLOR_TEXT)


def _add_bank_block(doc, amount: int, due_date_str: str) -> None:
    _para(doc, "", size_pt=4)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    _apply_font(p.add_run(f"銀行振込で ¥{amount:,} をお支払い"),
                size_pt=12, bold=True, color=COLOR_TEXT)
    _para(doc,
          f"お支払期限: {due_date_str}　※ 振込手数料はお客様ご負担にてお願いいたします。",
          size_pt=9, color=COLOR_MUTED, space_after_pt=6)

    bank_rows = [
        ("金融機関", config.BANK_NAME),
        ("支店名",   config.BANK_BRANCH),
        ("種別",     config.BANK_ACCOUNT_TYPE),
        ("口座番号", config.BANK_ACCOUNT_NUMBER),
        ("口座名義", config.BANK_ACCOUNT_HOLDER),
    ]
    tbl = doc.add_table(rows=len(bank_rows), cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(3.5); tbl.columns[1].width = Cm(13.5)
    for i, (k, v) in enumerate(bank_rows):
        c1 = tbl.rows[i].cells[0]; c1.width = Cm(3.5)
        c2 = tbl.rows[i].cells[1]; c2.width = Cm(13.5)
        for c in (c1, c2):
            _set_cell_border(c,
                top=_nil(), bottom=_single(sz="4"), left=_nil(), right=_nil())
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(3); p1.paragraph_format.space_after = Pt(3)
        _apply_font(p1.add_run(k), size_pt=9, color=COLOR_MUTED)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(3); p2.paragraph_format.space_after = Pt(3)
        _apply_font(p2.add_run(v), size_pt=10, color=COLOR_TEXT)


def _add_notes_block(doc, notes: list[str]) -> None:
    _para(doc, "", size_pt=6)
    _para(doc, "備考", size_pt=9, bold=True, color=COLOR_MUTED, space_after_pt=3)
    for n in notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.3)
        _apply_font(p.add_run(f"・{n}"), size_pt=9, color=COLOR_TEXT)


def _add_footer(doc) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _apply_font(p.add_run(f"{config.SELLER_NAME}　／　{config.SELLER_EMAIL}"),
                size_pt=8, color=COLOR_MUTED)


# ---- 日付フォーマット ----

def _jdate(d: datetime.date) -> str:
    return f"{d.year}年{d.month}月{d.day}日"


# ---- 書類ビルダー ----

def build_docx(params: DocumentParams, out_path: str) -> str:
    """Word 文書を生成して out_path に保存する。戻り値は out_path。"""
    _prepare_logo()
    doc = Document()
    _setup_page(doc)

    dt = params.doc_type

    if dt == DocType.QUOTE:
        title       = "見　積　書"
        amount_lbl  = "お見積金額"
        total_lbl   = "お見積合計"
        sub_note    = "下記明細のとおり、お見積りいたします。"
        meta = [
            ("見積書番号", params.doc_number),
            ("発行日　　", _jdate(params.issue_date)),
            ("有効期限　", _jdate(params.valid_until) if params.valid_until
                          else "発行日から 30 日以内"),
        ]
    elif dt == DocType.INVOICE:
        title       = "請　求　書"
        amount_lbl  = "ご請求金額"
        total_lbl   = "ご請求金額"
        sub_note    = "下記明細のとおり、ご請求申し上げます。"
        meta = [
            ("請求書番号", params.doc_number),
            ("発行日　　", _jdate(params.issue_date)),
            ("お支払期限", _jdate(params.due_date) if params.due_date else "別途ご連絡"),
        ]
    else:  # RECEIPT
        title       = "領　収　書"
        amount_lbl  = "領収金額"
        total_lbl   = "領収合計"
        sub_note    = "上記の金額を、正に領収いたしました。"
        meta = [
            ("領収書番号", params.doc_number),
            ("発行日　　", _jdate(params.issue_date)),
            ("領収日　　", _jdate(params.issue_date)),
        ]

    _add_header_block(doc, title, meta)
    _add_parties_block(doc, params.customer_name, params.customer_address)
    _add_headline_amount(doc,
        label=amount_lbl,
        amount_text=f"¥{params.total:,}（消費税不課税）",
        sub_note=sub_note)

    if dt == DocType.RECEIPT:
        _para(doc, "", size_pt=2)
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
        _apply_font(p.add_run("但し　"), size_pt=10, color=COLOR_MUTED)
        _apply_font(p.add_run(f"{params.items[0].description} 代金として"),
                    size_pt=11, bold=True, color=COLOR_TEXT)

    _add_line_items_table(doc, params.items)
    _add_totals_block(doc, params.total, total_lbl)

    if dt == DocType.INVOICE and params.due_date:
        _add_bank_block(doc, params.total, _jdate(params.due_date))

    if dt == DocType.QUOTE:
        notes = [
            "本見積書は有効期限内に限り有効です。",
            config.NOTE_LICENSE,
            "ご注文にあたっては Office Go Plan 利用規約にご同意いただく必要があります。",
            config.NOTE_TERMS_B2B,
            config.NOTE_TAX,
        ]
    elif dt == DocType.INVOICE:
        notes = [config.NOTE_LICENSE, config.NOTE_TERMS_B2B, config.NOTE_TAX]
    else:
        notes = [
            "本領収書は、入金確認後または決済承認後に発行しています。",
            config.NOTE_TAX,
        ]
    notes += params.notes_extra
    _add_notes_block(doc, notes)
    _add_footer(doc)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path


def to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Word COM 経由で .docx → .pdf 変換。成功時 True。"""
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            d = word.Documents.Open(os.path.abspath(docx_path))
            d.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            d.Close()
        finally:
            word.Quit()
        return True
    except Exception:
        return False


def generate(params: DocumentParams) -> dict[str, str | None]:
    """
    docx + pdf を生成する。
    戻り値: {"docx": path, "pdf": path_or_None}
    """
    subdir = {DocType.QUOTE: "quotes", DocType.INVOICE: "invoices",
              DocType.RECEIPT: "receipts"}[params.doc_type]
    out_dir   = os.path.join(config.OUTPUT_DIR, subdir)
    docx_path = os.path.join(out_dir, f"{params.doc_number}.docx")
    pdf_path  = os.path.join(out_dir, f"{params.doc_number}.pdf")

    build_docx(params, docx_path)
    result: dict[str, str | None] = {"docx": docx_path, "pdf": None}
    if to_pdf(docx_path, pdf_path):
        result["pdf"] = pdf_path
    return result
