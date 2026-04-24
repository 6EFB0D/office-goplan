# -*- coding: utf-8 -*-
"""
Office Go Plan - 見積書 / 請求書 / 領収書 Word テンプレート生成スクリプト

生成ファイル（本スクリプトと同じディレクトリに出力）:
  - template_quote.docx     (見積書)
  - template_invoice.docx   (請求書)
  - template_receipt.docx   (領収書)

設計方針:
  - Stripe が発行する請求書 PDF の体裁を踏襲し、3 種で統一感を確保。
  - 左上: 書類タイトル / 右上: ロゴ（assets/logo/logo-a.jpg）。
  - 大見出し金額 → 明細表 → 右寄せ合計 → 銀行振込ブロック（請求書のみ）。
  - 免税事業者（不課税）注記を各テンプレートに記載。
  - 法人後払いは利用規約 第5条の2 を参照する旨を付記（請求書のみ）。

利用方法:
  D:\\Users\\admin_mak\\project\\office-goplan\\docs\\invoice-templates> python generate_templates.py

依存:
  python-docx (>= 1.0)

最終更新: 2026-04-24
"""

from __future__ import annotations

import os
from copy import deepcopy

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


# ---------------------------------------------------------------------------
# 定数（事業者固定情報 / デザイン）
# ---------------------------------------------------------------------------

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_SOURCE_PATH = os.path.abspath(
    os.path.join(SCRIPT_DIR, "..", "..", "assets", "logo", "logo-a.jpg")
)
# python-docx は一部 JPEG のマーカー（例: APP11 等）を拒否するため、
# PIL で再エンコードした PNG をテンプレート埋め込み用に使用する。
LOGO_PATH = os.path.join(SCRIPT_DIR, "_logo_rendered.png")


def _prepare_logo() -> None:
    """ロゴ画像を python-docx で確実に読み込める PNG に再エンコード。"""
    if os.path.exists(LOGO_PATH):
        # 既に変換済みなら再利用（ソースより古い場合は上書き）
        try:
            src_mtime = os.path.getmtime(LOGO_SOURCE_PATH)
            dst_mtime = os.path.getmtime(LOGO_PATH)
            if dst_mtime >= src_mtime:
                return
        except OSError:
            pass
    if not os.path.exists(LOGO_SOURCE_PATH):
        return
    try:
        from PIL import Image  # 遅延 import（未導入環境では画像を諦めて続行）
    except ImportError:
        return
    with Image.open(LOGO_SOURCE_PATH) as img:
        img = img.convert("RGBA") if img.mode not in ("RGB", "RGBA") else img
        img.save(LOGO_PATH, format="PNG", optimize=True)

# 事業者情報（特商法表記・利用規約と整合させること）
SELLER_NAME = "Office Go Plan"
SELLER_POSTAL = "〒106-0032"
SELLER_ADDRESS = "東京都港区六本木2-1-19 S-Building 3F"
SELLER_EMAIL = "support@office-goplan.com"
SELLER_WEB = "https://office-goplan.com"

# 銀行振込先
# 注記: 現状は個人口座（屋号付きの場合あり）。法人口座へ切替時は本ブロックを更新すること。
BANK_NAME = "GMOあおぞらネット銀行"
BANK_BRANCH = "にじ支店（支店番号 302）"
BANK_ACCOUNT_TYPE = "普通"
BANK_ACCOUNT_NUMBER = "3189329"
# 口座名義（カタカナ、銀行登録の正式表記）
# 現状は屋号なしの個人名。屋号付き口座に切り替えた際はここを更新。
BANK_ACCOUNT_HOLDER = "カワギシ　マコト"

# カラー（ロゴの鮮緑を踏襲）
COLOR_ACCENT = RGBColor(0x00, 0xB8, 0x55)   # 見出し・ライン
COLOR_TEXT = RGBColor(0x1A, 0x1A, 0x1A)     # 本文
COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)    # 補足・ラベル
COLOR_DIVIDER = "D0D5DD"                    # 仕切り線（16進 RGB）

# フォント
FONT_JP = "Noto Sans JP"
FONT_FALLBACK = "Yu Gothic UI"


# ---------------------------------------------------------------------------
# OOXML ヘルパー
# ---------------------------------------------------------------------------

def _set_cell_border(cell, **edges) -> None:
    """セル罫線を設定。edges 例: top={'val': 'single', 'sz': '4', 'color': 'D0D5DD'}"""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        tag = qn(f"w:{edge_name}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge_name}")
            borders.append(element)
        for k, v in attrs.items():
            element.set(qn(f"w:{k}"), str(v))


def _shade_cell(cell, hex_color: str) -> None:
    """セル背景色を設定。hex_color は 'RRGGBB'。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _apply_font(run, name: str = FONT_JP, size_pt: float | None = None,
                bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(container, text: str = "", *, size_pt: float | None = None,
                   bold: bool = False, color: RGBColor | None = None,
                   align: int | None = None, space_after_pt: float | None = None,
                   line_spacing: float | None = None):
    p = container.add_paragraph() if hasattr(container, "add_paragraph") else container
    if align is not None:
        p.alignment = align
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        _apply_font(run, size_pt=size_pt, bold=bold, color=color)
    return p


# ---------------------------------------------------------------------------
# 各ブロック生成
# ---------------------------------------------------------------------------

def _setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # 既定スタイルの日本語フォント
    style = doc.styles["Normal"]
    style.font.name = FONT_JP
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_JP)
    rfonts.set(qn("w:hAnsi"), FONT_JP)
    rfonts.set(qn("w:eastAsia"), FONT_JP)


def _add_header_block(doc: Document, title: str, meta_rows: list[tuple[str, str]]) -> None:
    """左上タイトル＋右上ロゴ＋直下メタ情報ブロック。"""
    # 1列目にタイトル・メタ、2列目にロゴの 2x1 レイアウト
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(11.0)
    table.columns[1].width = Cm(6.0)

    left = table.rows[0].cells[0]
    left.width = Cm(11.0)
    right = table.rows[0].cells[1]
    right.width = Cm(6.0)

    # --- 左: タイトル ---
    title_p = left.paragraphs[0]
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run(title)
    _apply_font(run, size_pt=26, bold=True, color=COLOR_TEXT)

    # 緑アクセントのアンダーライン（細いテキストで表現）
    accent_p = left.add_paragraph()
    accent_p.paragraph_format.space_after = Pt(10)
    accent_run = accent_p.add_run("▬▬▬▬")
    _apply_font(accent_run, size_pt=10, color=COLOR_ACCENT, bold=True)

    # メタ情報（書類番号・発行日・期限など）
    for label, value in meta_rows:
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}　")
        _apply_font(r1, size_pt=9, color=COLOR_MUTED)
        r2 = p.add_run(value)
        _apply_font(r2, size_pt=10, color=COLOR_TEXT)

    # --- 右: ロゴ ---
    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(LOGO_PATH):
        run = right_p.add_run()
        run.add_picture(LOGO_PATH, width=Cm(5.0))
    else:
        run = right_p.add_run("[Office Go Plan ロゴ]")
        _apply_font(run, size_pt=10, color=COLOR_MUTED)

    # テーブル枠線なし
    for cell in (left, right):
        _set_cell_border(
            cell,
            top={"val": "nil"},
            bottom={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
        )


def _add_parties_block(doc: Document, buyer_name: str, buyer_address: str = "") -> None:
    """請求先（左）／発行元（右）ブロック。"""
    doc.add_paragraph()  # spacer

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.0)
    table.columns[1].width = Cm(8.0)

    # --- 左: 請求先 ---
    left = table.rows[0].cells[0]
    left.width = Cm(9.0)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("請求先")
    _apply_font(r, size_pt=9, color=COLOR_MUTED)

    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{buyer_name}　御中")
    _apply_font(r, size_pt=12, bold=True, color=COLOR_TEXT)

    if buyer_address:
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(buyer_address)
        _apply_font(r, size_pt=9, color=COLOR_TEXT)

    # --- 右: 発行元 ---
    right = table.rows[0].cells[1]
    right.width = Cm(8.0)
    p = right.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("発行元")
    _apply_font(r, size_pt=9, color=COLOR_MUTED)

    for txt, bold in [
        (SELLER_NAME, True),
        (f"{SELLER_POSTAL} {SELLER_ADDRESS}", False),
        (f"Email: {SELLER_EMAIL}", False),
        (SELLER_WEB, False),
    ]:
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt)
        _apply_font(r, size_pt=10 if bold else 9, bold=bold, color=COLOR_TEXT)

    for cell in (left, right):
        _set_cell_border(
            cell,
            top={"val": "nil"},
            bottom={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
        )


def _add_headline_amount(doc: Document, label: str, amount_text: str,
                         sub_note: str | None = None) -> None:
    """大見出しの金額ブロック。上下を薄いラインで区切る。"""
    # 上のライン
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    # スペーサー
    _add_paragraph(doc, "", size_pt=2)

    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_border(
        cell,
        top={"val": "single", "sz": "6", "color": COLOR_DIVIDER},
        bottom={"val": "single", "sz": "6", "color": COLOR_DIVIDER},
        left={"val": "nil"},
        right={"val": "nil"},
    )
    cell.paragraphs[0].paragraph_format.space_before = Pt(6)
    cell.paragraphs[0].paragraph_format.space_after = Pt(6)
    r = cell.paragraphs[0].add_run(f"{label}　")
    _apply_font(r, size_pt=11, color=COLOR_MUTED)
    r = cell.paragraphs[0].add_run(amount_text)
    _apply_font(r, size_pt=22, bold=True, color=COLOR_TEXT)
    if sub_note:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(sub_note)
        _apply_font(r, size_pt=9, color=COLOR_MUTED)


def _add_line_items_table(doc: Document, rows: list[tuple[str, int, int, int]]) -> None:
    """明細表。rows は [(説明, 数量, 単価, 金額), ...]。"""
    _add_paragraph(doc, "", size_pt=2)  # spacer

    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Cm(9.0), Cm(2.0), Cm(3.0), Cm(3.0)]
    for i, w in enumerate(widths):
        table.columns[i].width = w

    # ヘッダ行
    headers = ["説明", "数量", "単価", "金額"]
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].width = widths[i]
        _shade_cell(header_cells[i], "F6F8FA")
        _set_cell_border(
            header_cells[i],
            top={"val": "single", "sz": "6", "color": COLOR_DIVIDER},
            bottom={"val": "single", "sz": "6", "color": COLOR_DIVIDER},
            left={"val": "nil"},
            right={"val": "nil"},
        )
        p = header_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if i >= 1:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(h)
        _apply_font(r, size_pt=9, bold=True, color=COLOR_MUTED)

    # 明細行
    for idx, (desc, qty, unit, amount) in enumerate(rows, start=1):
        row = table.rows[idx].cells
        values = [desc, f"{qty:,}", f"￥{unit:,}", f"￥{amount:,}"]
        for i, v in enumerate(values):
            row[i].width = widths[i]
            _set_cell_border(
                row[i],
                top={"val": "nil"},
                bottom={"val": "single", "sz": "4", "color": COLOR_DIVIDER},
                left={"val": "nil"},
                right={"val": "nil"},
            )
            p = row[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            if i >= 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(v)
            _apply_font(r, size_pt=10, color=COLOR_TEXT)


def _add_totals_block(doc: Document, subtotal: int, total: int,
                      total_label: str = "請求金額") -> None:
    """右寄せ合計ブロック。"""
    _add_paragraph(doc, "", size_pt=2)

    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.columns[0].width = Cm(11.0)
    table.columns[1].width = Cm(6.0)

    lines = [
        ("小計", f"￥{subtotal:,}", False),
        ("合計", f"￥{total:,}", False),
        (total_label, f"￥{total:,}", True),
    ]
    for i, (lbl, val, emphasized) in enumerate(lines):
        left = table.rows[i].cells[0]
        right = table.rows[i].cells[1]
        left.width = Cm(11.0)
        right.width = Cm(6.0)

        # 空白セル（罫線なし）
        _set_cell_border(
            left,
            top={"val": "nil"}, bottom={"val": "nil"},
            left={"val": "nil"}, right={"val": "nil"},
        )

        # 値セル
        border_top = {"val": "single", "sz": "6", "color": COLOR_DIVIDER} if emphasized else {"val": "nil"}
        _set_cell_border(
            right,
            top=border_top,
            bottom={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
        )

        # 2 列表示（ラベル: 値）を 1 セル内で
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{lbl}　")
        _apply_font(r1, size_pt=10, color=COLOR_MUTED, bold=False)
        r2 = p.add_run(val)
        _apply_font(
            r2,
            size_pt=13 if emphasized else 10,
            bold=emphasized,
            color=COLOR_TEXT,
        )


def _add_bank_block(doc: Document, amount: int, due_date: str) -> None:
    """銀行振込のご案内ブロック。請求書のみ。"""
    _add_paragraph(doc, "", size_pt=4)
    _add_paragraph(
        doc,
        f"銀行振込で ￥{amount:,} をお支払い",
        size_pt=12,
        bold=True,
        color=COLOR_TEXT,
        space_after_pt=2,
    )
    _add_paragraph(
        doc,
        f"お支払期限: {due_date}　※ 振込手数料はお客様ご負担にてお願いいたします。",
        size_pt=9,
        color=COLOR_MUTED,
        space_after_pt=6,
    )

    rows = [
        ("金融機関", BANK_NAME),
        ("支店名", BANK_BRANCH),
        ("種別", BANK_ACCOUNT_TYPE),
        ("口座番号", BANK_ACCOUNT_NUMBER),
        ("口座名義", BANK_ACCOUNT_HOLDER),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(13.5)

    for i, (k, v) in enumerate(rows):
        c1 = table.rows[i].cells[0]
        c2 = table.rows[i].cells[1]
        c1.width = Cm(3.5)
        c2.width = Cm(13.5)
        for c in (c1, c2):
            _set_cell_border(
                c,
                top={"val": "nil"},
                bottom={"val": "single", "sz": "4", "color": COLOR_DIVIDER},
                left={"val": "nil"},
                right={"val": "nil"},
            )
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        r = p1.add_run(k)
        _apply_font(r, size_pt=9, color=COLOR_MUTED)

        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(3)
        p2.paragraph_format.space_after = Pt(3)
        r = p2.add_run(v)
        _apply_font(r, size_pt=10, color=COLOR_TEXT)


def _add_notes_block(doc: Document, notes: list[str]) -> None:
    """備考・注記ブロック（免税事業者・利用規約参照など）。"""
    _add_paragraph(doc, "", size_pt=6)
    _add_paragraph(
        doc, "備考",
        size_pt=9, bold=True, color=COLOR_MUTED, space_after_pt=3,
    )
    for n in notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(f"・{n}")
        _apply_font(r, size_pt=9, color=COLOR_TEXT)


def _add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{SELLER_NAME}　／　{SELLER_EMAIL}")
    _apply_font(r, size_pt=8, color=COLOR_MUTED)


# ---------------------------------------------------------------------------
# 共通備考
# ---------------------------------------------------------------------------

NOTE_TAX_EXEMPT = (
    "当サービスの運営者は、発行日現在において免税事業者であり、"
    "適格請求書発行事業者の登録番号（T番号）を有していません。"
    "表示金額は「不課税（免税事業者）」として税込一本価格で表示しており、"
    "消費税額は別建てで表示・請求しておりません。"
    "仕入税額控除の経過措置（〜2026年9月30日は相当額の80%、"
    "2026年10月1日〜2029年9月30日は50%）をご参照ください。"
)

NOTE_TERMS_B2B = (
    "法人ユーザーの後払い取引の条件（催告・一時停止・失効・遅延損害金 年14.6% 等）は、"
    "利用規約 第5条の2 に定めます（https://office-goplan.com/terms-of-service.html）。"
)

NOTE_LICENSE_DELIVERY = (
    "ライセンスキーは、Stripe 決済の場合は決済承認後、"
    "銀行振込（後払いを含む）の場合は入金確認後に、"
    "ご注文時にご指定のメールアドレス宛にお送りします。"
)


# ---------------------------------------------------------------------------
# テンプレート: 見積書
# ---------------------------------------------------------------------------

def build_quote() -> Document:
    doc = Document()
    _setup_page(doc)

    _add_header_block(
        doc,
        title="見　積　書",
        meta_rows=[
            ("見積書番号", "Q-YYYYMMDD-0001"),
            ("発行日　　", "YYYY年M月D日"),
            ("有効期限　", "発行日から 30 日以内"),
        ],
    )

    _add_parties_block(
        doc,
        buyer_name="株式会社〇〇〇〇",
        buyer_address="〒000-0000 東京都〇〇区〇〇町0-0-0",
    )

    _add_headline_amount(
        doc,
        label="お見積金額",
        amount_text="￥0,000（税込・不課税）",
        sub_note="下記明細のとおり、お見積りいたします。",
    )

    _add_line_items_table(
        doc,
        rows=[
            ("PDF Handler 買い切り版ライセンス", 1, 3850, 3850),
            # 複数行の例（削除可）:
            # ("ZipSearch 買い切り版ライセンス", 1, 0, 0),
        ],
    )

    _add_totals_block(doc, subtotal=3850, total=3850, total_label="お見積合計")

    _add_notes_block(
        doc,
        notes=[
            "本見積書は、発行日から 30 日以内に限り有効です。",
            NOTE_LICENSE_DELIVERY,
            "ご注文にあたっては、Office Go Plan 利用規約にご同意いただく必要があります。",
            NOTE_TERMS_B2B,
            NOTE_TAX_EXEMPT,
        ],
    )

    _add_footer(doc)
    return doc


# ---------------------------------------------------------------------------
# テンプレート: 請求書
# ---------------------------------------------------------------------------

def build_invoice() -> Document:
    doc = Document()
    _setup_page(doc)

    _add_header_block(
        doc,
        title="請　求　書",
        meta_rows=[
            ("請求書番号", "I-YYYYMMDD-0001"),
            ("発行日　　", "YYYY年M月D日"),
            ("お支払期限", "YYYY年M月D日"),
        ],
    )

    _add_parties_block(
        doc,
        buyer_name="株式会社〇〇〇〇",
        buyer_address="〒000-0000 東京都〇〇区〇〇町0-0-0",
    )

    _add_headline_amount(
        doc,
        label="ご請求金額",
        amount_text="￥0,000（税込・不課税）",
        sub_note="下記明細のとおり、ご請求申し上げます。",
    )

    _add_line_items_table(
        doc,
        rows=[
            ("PDF Handler 買い切り版ライセンス", 1, 3850, 3850),
        ],
    )

    _add_totals_block(doc, subtotal=3850, total=3850, total_label="ご請求金額")

    _add_bank_block(doc, amount=3850, due_date="YYYY年M月D日")

    _add_notes_block(
        doc,
        notes=[
            NOTE_LICENSE_DELIVERY,
            NOTE_TERMS_B2B,
            NOTE_TAX_EXEMPT,
        ],
    )

    _add_footer(doc)
    return doc


# ---------------------------------------------------------------------------
# テンプレート: 領収書
# ---------------------------------------------------------------------------

def build_receipt() -> Document:
    doc = Document()
    _setup_page(doc)

    _add_header_block(
        doc,
        title="領　収　書",
        meta_rows=[
            ("領収書番号", "R-YYYYMMDD-0001"),
            ("発行日　　", "YYYY年M月D日"),
            ("領収日　　", "YYYY年M月D日"),
        ],
    )

    _add_parties_block(
        doc,
        buyer_name="株式会社〇〇〇〇",
        buyer_address="〒000-0000 東京都〇〇区〇〇町0-0-0",
    )

    _add_headline_amount(
        doc,
        label="領収金額",
        amount_text="￥0,000（税込・不課税）",
        sub_note="上記の金額を、正に領収いたしました。",
    )

    # 但し書き（領収書の特徴）
    _add_paragraph(doc, "", size_pt=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("但し　")
    _apply_font(r, size_pt=10, color=COLOR_MUTED)
    r = p.add_run("PDF Handler 買い切り版ライセンス 代金として")
    _apply_font(r, size_pt=11, bold=True, color=COLOR_TEXT)

    _add_line_items_table(
        doc,
        rows=[
            ("PDF Handler 買い切り版ライセンス", 1, 3850, 3850),
        ],
    )

    _add_totals_block(doc, subtotal=3850, total=3850, total_label="領収合計")

    _add_notes_block(
        doc,
        notes=[
            "本領収書は、銀行振込による入金確認後、または Stripe 等決済承認後に発行しています。",
            NOTE_TAX_EXEMPT,
            "※ 本領収書は収入印紙が必要となる金額（5万円以上）に該当する場合、"
            "法令に従って印紙を貼付し消印のうえお渡しします（電子発行の場合は課税文書に該当しないため不要）。",
        ],
    )

    _add_footer(doc)
    return doc


# ---------------------------------------------------------------------------
# エントリポイント
# ---------------------------------------------------------------------------

def main() -> None:
    _prepare_logo()
    out_dir = SCRIPT_DIR
    jobs = [
        ("template_quote.docx", build_quote),
        ("template_invoice.docx", build_invoice),
        ("template_receipt.docx", build_receipt),
    ]
    for fname, builder in jobs:
        doc = builder()
        path = os.path.join(out_dir, fname)
        doc.save(path)
        print(f"  generated: {path}")

    print("\n完了: 3 ファイルを生成しました。")
    print("  - 見積書 (template_quote.docx)")
    print("  - 請求書 (template_invoice.docx)")
    print("  - 領収書 (template_receipt.docx)")


if __name__ == "__main__":
    main()
